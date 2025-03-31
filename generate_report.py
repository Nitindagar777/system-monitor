from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from datetime import datetime

def create_report():
    doc = Document()
    
    # Set basic styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('System Monitoring Application', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('A Comprehensive Real-time System Resource Monitoring Tool')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author = doc.add_paragraph('Prepared by:\nNitin')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %Y")}')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    summary = doc.add_paragraph()
    summary.add_run('''The System Monitoring Application represents a significant advancement in system resource monitoring technology. This comprehensive solution provides real-time insights into system performance through an intuitive interface while maintaining minimal resource overhead.

Key Features and Achievements:
• Real-time system metrics monitoring with millisecond precision
• Comprehensive process management and tracking
• Advanced network traffic analysis and visualization
• Detailed disk usage monitoring and analytics
• Sophisticated log management system
• Intuitive dark-themed user interface
• Efficient resource utilization
• Cross-platform compatibility

The application successfully addresses the challenges of modern system monitoring:
1. Real-time Performance Tracking
2. Resource Usage Optimization
3. System Health Monitoring
4. Network Traffic Analysis
5. Storage Management
6. Process Control and Monitoring

Technical Implementation Highlights:
• Backend: Python with Flask framework
• Frontend: HTML5, CSS3, JavaScript
• Real-time Updates: WebSocket implementation
• Data Collection: psutil library integration
• Database: SQLite for data persistence
• Authentication: JWT-based security
''')
    
    doc.add_page_break()
    
    # Technology Stack
    doc.add_heading('Technology Stack', 1)
    tech_stack = doc.add_paragraph()
    tech_stack.add_run('''The application utilizes a modern technology stack designed for optimal performance and maintainability:

1. Backend Technologies:
   • Python 3.9+
     - Primary programming language
     - Extensive standard library
     - Strong community support
     - Cross-platform compatibility
   
   • Flask Framework
     - Lightweight web framework
     - RESTful API support
     - Extension ecosystem
     - Easy integration
   
   • psutil Library
     - System information retrieval
     - Cross-platform support
     - Comprehensive metrics collection
     - Efficient resource usage

2. Frontend Technologies:
   • HTML5
     - Semantic markup
     - Modern web standards
     - Cross-browser compatibility
   
   • CSS3
     - Responsive design
     - Dark theme implementation
     - Custom animations
     - Flexible layouts
   
   • JavaScript
     - Real-time updates
     - Dynamic UI elements
     - Chart rendering
     - WebSocket handling

3. Database:
   • SQLite
     - Lightweight database
     - Zero configuration
     - ACID compliance
     - Built-in support

4. Additional Libraries:
   • Chart.js
     - Interactive charts
     - Real-time updates
     - Customizable themes
   
   • Socket.IO
     - WebSocket implementation
     - Real-time communication
     - Fallback support
''')
    
    doc.add_page_break()
    
    # Core Features
    doc.add_heading('Core Features', 1)
    
    # CPU Monitoring
    doc.add_heading('CPU Monitoring', 2)
    cpu_section = doc.add_paragraph()
    cpu_section.add_run('''The CPU monitoring module provides comprehensive insights into processor utilization and performance:

1. Real-time Metrics:
   • Overall CPU usage
   • Per-core utilization
   • Frequency scaling
   • Temperature monitoring
   • Load averages

2. Implementation Details:
   ```python
   class CPUMonitor:
       def __init__(self):
           self.history = deque(maxlen=3600)
           self.update_interval = 0.1
           
       def get_metrics(self):
           cpu_freq = psutil.cpu_freq()
           cpu_percent = psutil.cpu_percent(interval=0.1, percpu=True)
           
           return {
               'usage': cpu_percent,
               'frequency': cpu_freq.current,
               'cores': psutil.cpu_count(),
               'load_avg': psutil.getloadavg()
           }
   ```

3. Visualization:
   • Real-time usage graphs
   • Per-core usage bars
   • Temperature indicators
   • Historical trends
''')
    
    # Memory Monitoring
    doc.add_heading('Memory Monitoring', 2)
    memory_section = doc.add_paragraph()
    memory_section.add_run('''The memory monitoring system tracks RAM and virtual memory usage:

1. Key Metrics:
   • Physical memory usage
   • Virtual memory statistics
   • Swap space utilization
   • Page file statistics
   • Cache usage

2. Implementation:
   ```python
   class MemoryMonitor:
       def __init__(self):
           self.history = deque(maxlen=3600)
           
       def get_memory_info(self):
           memory = psutil.virtual_memory()
           swap = psutil.swap_memory()
           
           return {
               'total': memory.total,
               'available': memory.available,
               'used': memory.used,
               'free': memory.free,
               'swap_total': swap.total,
               'swap_used': swap.used
           }
   ```

3. Features:
   • Real-time memory tracking
   • Usage threshold alerts
   • Memory leak detection
   • Swap usage monitoring
''')
    
    doc.add_page_break()
    
    # Network Monitoring
    doc.add_heading('Network Monitoring', 2)
    network_section = doc.add_paragraph()
    network_section.add_run('''The network monitoring module provides detailed insights into network activity and performance:

1. Core Functionality:
   • Interface statistics tracking
   • Bandwidth monitoring
   • Connection tracking
   • Protocol analysis
   • Network error detection

2. Implementation Details:
   ```python
   class NetworkMonitor:
       def __init__(self):
           self.previous_counters = None
           self.current_counters = None
           self.update_interval = 1.0
           
       def get_network_stats(self):
           interfaces = psutil.net_if_stats()
           io_counters = psutil.net_io_counters(pernic=True)
           connections = psutil.net_connections()
           
           stats = {
               'interfaces': {},
               'connections': [],
               'io_counters': {}
           }
           
           # Process interface statistics
           for interface, stats in interfaces.items():
               stats['interfaces'][interface] = {
                   'isup': stats.isup,
                   'speed': stats.speed,
                   'mtu': stats.mtu,
                   'duplex': stats.duplex
               }
           
           # Process IO counters
           for nic, counters in io_counters.items():
               stats['io_counters'][nic] = {
                   'bytes_sent': counters.bytes_sent,
                   'bytes_recv': counters.bytes_recv,
                   'packets_sent': counters.packets_sent,
                   'packets_recv': counters.packets_recv,
                   'errin': counters.errin,
                   'errout': counters.errout,
                   'dropin': counters.dropin,
                   'dropout': counters.dropout
               }
           
           return stats
   ```

3. Features:
   • Real-time bandwidth monitoring
   • Network interface status
   • Connection tracking
   • Error rate monitoring
   • Protocol distribution analysis

4. Visualization:
   • Bandwidth usage graphs
   • Connection tables
   • Protocol distribution charts
   • Error rate indicators
''')

    # Disk Monitoring
    doc.add_heading('Disk Monitoring', 2)
    disk_section = doc.add_paragraph()
    disk_section.add_run('''The disk monitoring system provides comprehensive storage analytics:

1. Key Features:
   • Partition information
   • Usage statistics
   • I/O monitoring
   • Performance metrics
   • SMART status (where available)

2. Implementation:
   ```python
   class DiskMonitor:
       def __init__(self):
           self.disk_history = {}
           self.update_interval = 2.0
           
       def get_disk_info(self):
           partitions = psutil.disk_partitions()
           io_counters = psutil.disk_io_counters(perdisk=True)
           
           disk_info = {
               'partitions': {},
               'io_stats': {}
           }
           
           for partition in partitions:
               try:
                   usage = psutil.disk_usage(partition.mountpoint)
                   disk_info['partitions'][partition.device] = {
                       'mountpoint': partition.mountpoint,
                       'fstype': partition.fstype,
                       'total': usage.total,
                       'used': usage.used,
                       'free': usage.free,
                       'percent': usage.percent
                   }
               except PermissionError:
                   continue
           
           for disk, counters in io_counters.items():
               disk_info['io_stats'][disk] = {
                   'read_count': counters.read_count,
                   'write_count': counters.write_count,
                   'read_bytes': counters.read_bytes,
                   'write_bytes': counters.write_bytes,
                   'read_time': counters.read_time,
                   'write_time': counters.write_time
               }
           
           return disk_info
   ```

3. Monitoring Capabilities:
   • Storage space tracking
   • I/O performance analysis
   • Disk health monitoring
   • Usage trend analysis
   • Alert system for space issues

4. Visual Elements:
   • Usage pie charts
   • I/O activity graphs
   • Storage trend lines
   • Health status indicators
''')

    # Process Management
    doc.add_heading('Process Management', 2)
    process_section = doc.add_paragraph()
    process_section.add_run('''The process management module provides detailed information about running processes:

1. Core Features:
   • Process listing
   • Resource usage tracking
   • Process control
   • Thread management
   • Performance monitoring

2. Implementation:
   ```python
   class ProcessManager:
       def __init__(self):
           self.process_list = {}
           self.update_interval = 0.5
           
       def get_process_list(self):
           processes = []
           for proc in psutil.process_iter(['pid', 'name', 'username', 'cpu_percent', 'memory_percent', 'status']):
               try:
                   pinfo = proc.info
                   pinfo.update({
                       'threads': proc.num_threads(),
                       'create_time': datetime.fromtimestamp(proc.create_time()).strftime('%Y-%m-%d %H:%M:%S'),
                       'memory_info': proc.memory_info()._asdict(),
                       'io_counters': proc.io_counters()._asdict() if hasattr(proc, 'io_counters') else None
                   })
                   processes.append(pinfo)
               except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                   continue
           return processes
           
       def get_process_details(self, pid):
           try:
               proc = psutil.Process(pid)
               return {
                   'pid': proc.pid,
                   'name': proc.name(),
                   'status': proc.status(),
                   'cpu_percent': proc.cpu_percent(),
                   'memory_percent': proc.memory_percent(),
                   'threads': proc.num_threads(),
                   'connections': proc.connections(),
                   'open_files': proc.open_files(),
                   'memory_maps': proc.memory_maps(),
                   'environ': proc.environ()
               }
           except psutil.NoSuchProcess:
               return None
   ```

3. Process Information:
   • Basic process details
   • Resource consumption
   • Thread information
   • File handles
   • Network connections

4. Management Features:
   • Process filtering
   • Resource usage alerts
   • Process termination
   • Priority adjustment
   • Memory analysis

5. Visualization:
   • Process list table
   • Resource usage graphs
   • Thread count charts
   • Memory maps
   • Connection diagrams
''')

    doc.add_page_break()

    # System Logs
    doc.add_heading('System Logs Management', 2)
    logs_section = doc.add_paragraph()
    logs_section.add_run('''The system logs management module provides comprehensive log monitoring and analysis:

1. Key Features:
   • Real-time log monitoring
   • Log file management
   • Error tracking
   • Pattern recognition
   • Alert system

2. Implementation:
   ```python
   class LogManager:
       def __init__(self):
           self.log_paths = {
               'system': '/var/log/syslog',
               'auth': '/var/log/auth.log',
               'application': '/var/log/application.log'
           }
           self.patterns = {
               'error': r'ERROR|CRITICAL|FATAL',
               'warning': r'WARNING|WARN',
               'info': r'INFO|NOTICE'
           }
           
       def monitor_logs(self, log_type):
           if log_type not in self.log_paths:
               return None
               
           log_entries = []
           try:
               with open(self.log_paths[log_type], 'r') as f:
                   for line in f.readlines()[-100:]:
                       entry = self.parse_log_entry(line)
                       if entry:
                           log_entries.append(entry)
           except Exception as e:
               print(f"Error reading log file: {str(e)}")
               
           return log_entries
           
       def parse_log_entry(self, line):
           # Implementation of log parsing logic
           pass
   ```

3. Log Analysis Features:
   • Pattern matching
   • Error categorization
   • Trend analysis
   • Statistical reporting
   • Alert generation

4. Visualization:
   • Log entry table
   • Error rate graphs
   • Pattern distribution charts
   • Timeline views
   • Alert notifications
''')

    # User Interface Design
    doc.add_heading('User Interface Design', 1)
    ui_section = doc.add_paragraph()
    ui_section.add_run('''The user interface is designed with modern web standards and best practices:

1. Design Philosophy:
   • Clean and minimalist approach
   • Dark theme for reduced eye strain
   • Responsive layout for all screen sizes
   • Intuitive navigation
   • Consistent styling

2. Component Layout:
   • Sidebar navigation
   • Main content area
   • Status bar
   • Modal dialogs
   • Toast notifications

3. Implementation:
   ```html
   <!DOCTYPE html>
   <html lang="en">
   <head>
       <meta charset="UTF-8">
       <meta name="viewport" content="width=device-width, initial-scale=1.0">
       <title>System Monitor</title>
       <link rel="stylesheet" href="static/css/style.css">
   </head>
   <body class="dark-theme">
       <div class="sidebar">
           <nav class="nav-menu">
               <ul>
                   <li><a href="/">Dashboard</a></li>
                   <li><a href="/processes">Processes</a></li>
                   <li><a href="/network">Network</a></li>
                   <li><a href="/disks">Storage</a></li>
                   <li><a href="/logs">Logs</a></li>
               </ul>
           </nav>
       </div>
       
       <main class="content">
           <div class="metrics-grid">
               <div class="metric-card">
                   <h3>CPU Usage</h3>
                   <div class="chart-container">
                       <canvas id="cpuChart"></canvas>
                   </div>
               </div>
               
               <div class="metric-card">
                   <h3>Memory Usage</h3>
                   <div class="chart-container">
                       <canvas id="memoryChart"></canvas>
                   </div>
               </div>
           </div>
       </main>
   </body>
   </html>
   ```

4. CSS Implementation:
   ```css
   /* Dark theme variables */
   :root {
       --bg-primary: #1a1a1a;
       --bg-secondary: #2d2d2d;
       --text-primary: #ffffff;
       --text-secondary: #b3b3b3;
       --accent-color: #007acc;
   }

   body.dark-theme {
       background-color: var(--bg-primary);
       color: var(--text-primary);
   }

   .sidebar {
       background-color: var(--bg-secondary);
       width: 250px;
       height: 100vh;
       position: fixed;
   }

   .metric-card {
       background-color: var(--bg-secondary);
       border-radius: 8px;
       padding: 1rem;
       margin: 1rem;
   }

   .chart-container {
       height: 300px;
       position: relative;
   }
   ```

5. JavaScript Implementation:
   ```javascript
   class DashboardUI {
       constructor() {
           this.charts = {};
           this.updateInterval = 1000;
           this.init();
       }
       
       async init() {
           this.setupCharts();
           this.setupWebSocket();
           this.startUpdates();
       }
       
       setupCharts() {
           this.charts.cpu = new Chart('cpuChart', {
               type: 'line',
               data: {
                   labels: [],
                   datasets: [{
                       label: 'CPU Usage',
                       data: [],
                       borderColor: '#007acc'
                   }]
               },
               options: {
                   responsive: true,
                   animation: { duration: 0 }
               }
           });
           
           // Similar setup for other charts
       }
       
       async updateMetrics() {
           try {
               const response = await fetch('/api/metrics');
               const data = await response.json();
               
               this.updateCharts(data);
               this.updateCards(data);
           } catch (error) {
               console.error('Error updating metrics:', error);
           }
       }
   }
   ```''')

    # Performance Optimization
    doc.add_heading('Performance Optimization', 1)
    performance_section = doc.add_paragraph()
    performance_section.add_run('''The application implements various optimization techniques to ensure efficient operation:

1. Backend Optimization:
   • Efficient data collection
   • Caching mechanisms
   • Database query optimization
   • Memory management
   • Resource pooling

2. Frontend Optimization:
   • Minimized asset sizes
   • Efficient DOM updates
   • Debounced event handlers
   • Lazy loading
   • Virtual scrolling

3. Network Optimization:
   • WebSocket for real-time updates
   • Compressed data transfer
   • Request batching
   • Connection pooling
   • Cache headers

4. Implementation Examples:
   ```python
   class PerformanceOptimizer:
       def __init__(self):
           self.cache = {}
           self.cache_timeout = 60  # seconds
           
       def get_cached_data(self, key):
           if key in self.cache:
               data, timestamp = self.cache[key]
               if time.time() - timestamp < self.cache_timeout:
                   return data
           return None
           
       def set_cached_data(self, key, data):
           self.cache[key] = (data, time.time())
           
       def clean_old_cache(self):
           current_time = time.time()
           self.cache = {
               k: v for k, v in self.cache.items()
               if current_time - v[1] < self.cache_timeout
           }
   ```

5. Database Optimization:
   ```python
   class DatabaseOptimizer:
       def __init__(self):
           self.connection_pool = []
           self.max_connections = 10
           
       def get_connection(self):
           if not self.connection_pool:
               return self.create_connection()
           return self.connection_pool.pop()
           
       def release_connection(self, connection):
           if len(self.connection_pool) < self.max_connections:
               self.connection_pool.append(connection)
           else:
               connection.close()
   ```''')

    # Testing and Quality Assurance
    doc.add_heading('Testing and Quality Assurance', 1)
    testing_section = doc.add_paragraph()
    testing_section.add_run('''The application undergoes rigorous testing to ensure reliability:

1. Unit Testing:
   ```python
   class TestCPUMonitor(unittest.TestCase):
       def setUp(self):
           self.monitor = CPUMonitor()
           
       def test_cpu_metrics(self):
           metrics = self.monitor.get_metrics()
           self.assertIsNotNone(metrics)
           self.assertIn('usage', metrics)
           self.assertIn('frequency', metrics)
           
       def test_historical_data(self):
           data = self.monitor.get_historical_data()
           self.assertIsInstance(data, list)
           
   class TestMemoryMonitor(unittest.TestCase):
       def setUp(self):
           self.monitor = MemoryMonitor()
           
       def test_memory_info(self):
           info = self.monitor.get_memory_info()
           self.assertIn('total', info)
           self.assertIn('available', info)
           self.assertTrue(info['total'] > 0)
   ```

2. Integration Testing:
   ```python
   class TestSystemMonitor(unittest.TestCase):
       def setUp(self):
           self.app = create_app('testing')
           self.client = self.app.test_client()
           
       def test_metrics_endpoint(self):
           response = self.client.get('/api/metrics')
           self.assertEqual(response.status_code, 200)
           data = json.loads(response.data)
           self.assertIn('cpu', data)
           self.assertIn('memory', data)
           
       def test_process_list(self):
           response = self.client.get('/api/processes')
           self.assertEqual(response.status_code, 200)
           processes = json.loads(response.data)
           self.assertIsInstance(processes, list)
   ```

3. Performance Testing:
   ```python
   class PerformanceTests:
       def __init__(self):
           self.monitor = SystemMonitor()
           
       def test_response_time(self):
           start_time = time.time()
           metrics = self.monitor.get_system_metrics()
           end_time = time.time()
           
           response_time = end_time - start_time
           assert response_time < 0.1, f"Response time too high: {response_time}s"
           
       def test_memory_usage(self):
           process = psutil.Process()
           initial_memory = process.memory_info().rss
           
           # Perform operations
           self.monitor.get_system_metrics()
           
           final_memory = process.memory_info().rss
           memory_increase = final_memory - initial_memory
           
           assert memory_increase < 10 * 1024 * 1024, "Memory usage too high"
   ```''')

    # Future Enhancements
    doc.add_heading('Future Enhancements', 1)
    future_section = doc.add_paragraph()
    future_section.add_run('''Planned improvements and future features:

1. Technical Enhancements:
   • Machine learning for anomaly detection
   • Predictive analytics for resource usage
   • Advanced log analysis
   • Custom metric creation
   • Extended API capabilities

2. User Interface Improvements:
   • Customizable dashboards
   • Additional visualization options
   • Mobile application
   • Desktop notifications
   • Keyboard shortcuts

3. Feature Additions:
   • Remote system monitoring
   • Cluster monitoring
   • Container monitoring
   • Cloud integration
   • Automated reporting

4. Performance Improvements:
   • Enhanced caching
   • Optimized data collection
   • Reduced resource usage
   • Faster page loads
   • Better compression''')

    # Conclusion
    doc.add_heading('Conclusion', 1)
    conclusion = doc.add_paragraph()
    conclusion.add_run('''The System Monitoring Application represents a significant achievement in system resource monitoring and management. Through its comprehensive feature set, efficient implementation, and user-friendly interface, it provides administrators and users with powerful tools for system analysis and optimization.

Key Accomplishments:
1. Real-time Monitoring
   • Accurate system metrics
   • Minimal resource overhead
   • Comprehensive data collection
   • Efficient data presentation

2. User Experience
   • Intuitive interface
   • Responsive design
   • Dark theme
   • Easy navigation

3. Technical Achievement
   • Efficient implementation
   • Robust architecture
   • Extensive testing
   • Quality assurance

4. Future Potential
   • Expandable framework
   • Modular design
   • API integration
   • Feature extensibility

The application successfully meets its objectives of providing comprehensive system monitoring while maintaining efficiency and usability. Its modular design and extensive documentation ensure easy maintenance and future development.''')

    # Save the document
    doc.save('detailed_report.docx')

if __name__ == '__main__':
    create_report() 